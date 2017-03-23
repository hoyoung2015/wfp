from docx import Document

from docx.shared import Inches


f = '/media/hoyoung/win7 home/compage/002043/01b75fb5169e58b746358d66e12d05f6a3ee36d6.docx'
f = '/home/hoyoung/01b75fb5169e58b746358d66e12d05f6a3ee36d6.doc'
doc = Document(f)
print(len(doc.paragraphs))
print(doc.paragraphs[17].text)